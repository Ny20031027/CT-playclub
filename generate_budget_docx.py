# -*- coding: utf-8 -*-
"""生成 PlayClub 运营成本预算报告 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
import os

def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(2.54)
    section.left_margin = section.right_margin = Cm(2.54)
    section.orientation = WD_ORIENTATION.PORTRAIT

def tune_styles(doc):
    # 正文样式
    body = doc.styles['Normal']
    body.font.name = '微软雅黑'
    body.font.size = Pt(10.5)
    body.paragraph_format.line_spacing = 1.5
    body.paragraph_format.space_after = Pt(6)
    body.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题样式
    title = doc.styles['Title']
    title.font.name = '微软雅黑'
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    title.paragraph_format.space_after = Pt(12)
    title.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 副标题样式
    subtitle = doc.styles['Subtitle']
    subtitle.font.name = '微软雅黑'
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    subtitle.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题1样式
    h1 = doc.styles['Heading 1']
    h1.font.name = '微软雅黑'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    h1.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题2样式
    h2 = doc.styles['Heading 2']
    h2.font.name = '微软雅黑'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = '微软雅黑'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    return table

def add_bold_text(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def main():
    doc = Document()
    setup_page(doc)
    tune_styles(doc)

    # 封面
    doc.add_paragraph('PlayClub（CT电竞陪玩平台）', style='Title')
    doc.add_paragraph('云托管运营成本预算报告', style='Subtitle')
    doc.add_paragraph()
    doc.add_paragraph('基于微信云托管官方定价 + 腾讯云 COS 实际配置')
    doc.add_paragraph('编制日期：2026年7月19日')
    doc.add_paragraph('编制单位：PlayClub 技术团队')
    doc.add_paragraph()
    doc.add_paragraph('仅供投资决策参考', style='Subtitle')
    doc.add_page_break()

    # 一、项目概述
    doc.add_heading('一、项目概述', level=1)
    doc.add_paragraph('PlayClub 是一个游戏代练/陪玩服务平台，技术架构如下：')
    add_table(doc, ['组件', '技术栈'], [
        ['后端', 'Django 4.1 + Django REST Framework'],
        ['数据库', '微信云托管 MySQL 5.7（Serverless）'],
        ['文件存储', '腾讯云 COS 对象存储'],
        ['前端', 'Vue.js 管理后台 + 微信小程序'],
        ['部署', '微信云托管（Docker 容器）'],
    ])
    doc.add_paragraph('核心功能：用户管理、订单系统、排班系统、财务统计、陪玩师列表、消息通知。')

    # 二、用户行为模型
    doc.add_heading('二、用户行为模型（1000 DAU）', level=1)
    add_table(doc, ['指标', '数值', '说明'], [
        ['日活用户（DAU）', '1,000', '目标规模'],
        ['人均日 API 请求', '30 次', '首页、列表、详情、订单、个人中心'],
        ['日总请求量', '30,000 次', '1000 × 30'],
        ['活跃时段', '08:00-24:00', '16 小时'],
        ['高峰时段', '12:00-14:00, 19:00-23:00', '约 6 小时'],
        ['日均文件上传', '200 次', '头像、轮播图、订单图片'],
        ['单文件平均大小', '200 KB', '图片为主'],
    ])

    # 三、资源配置方案
    doc.add_heading('三、资源配置方案', level=1)
    add_table(doc, ['资源', '配置', '说明'], [
        ['容器（服务）', '1 核 2GB', 'Django + Gunicorn，3 个 worker'],
        ['实例副本数', '最小 1，最大 3', '保留 1 个实例避免冷启动'],
        ['MySQL', '0.25-0.5 CCU 弹性', 'Serverless 自动伸缩'],
        ['数据库存储', '5GB', '用户、订单、排班等数据'],
        ['COS 对象存储', '5GB', '头像、轮播图、订单图片'],
        ['COS CDN', '公网访问', '图片加速分发'],
    ])

    # 四、微信云托管计费规则
    doc.add_heading('四、微信云托管计费规则', level=1)
    add_table(doc, ['资源类型', '计费项', '单价'], [
        ['容器', 'CPU', '0.055 元/(核·小时)'],
        ['容器', '内存', '0.032 元/(GB·小时)'],
        ['容器', '公网流量', '0.8 元/GB'],
        ['MySQL', '算力', '0.342 元/(CCU·小时)'],
        ['MySQL', '存储', '0.00485 元/(GB·小时)'],
        ['对象存储', '容量', '0.0043 元/(GB·天)'],
        ['对象存储', '下载次数', '0.01 元/万次'],
        ['对象存储', 'CDN 流量', '0.18 元/GB'],
        ['构建', '构建时长', '0.05 元/分钟'],
    ])

    doc.add_heading('免费额度（首 3 个月）', level=2)
    add_table(doc, ['资源', '免费额度'], [
        ['容器 CPU', '720 核·时'],
        ['容器内存', '1440 GB·时'],
        ['MySQL 算力', '720 CCU·时'],
        ['公网流量', '5 GB'],
    ])

    # 五、详细成本测算
    doc.add_heading('五、详细成本测算', level=1)

    doc.add_heading('5.1 容器（服务）费用', level=2)
    add_table(doc, ['时段', '时长', '配置', 'CPU 消耗', '内存消耗'], [
        ['高峰', '6h', '1核2GB × 2实例', '12 核·时', '24 GB·时'],
        ['正常', '9h', '1核2GB × 1实例', '9 核·时', '18 GB·时'],
        ['低谷', '1h', '1核2GB × 1实例', '1 核·时', '2 GB·时'],
        ['深夜', '8h', '保留 1 实例', '8 核·时', '16 GB·时'],
        ['日合计', '', '', '30 核·时', '60 GB·时'],
    ])
    add_table(doc, ['周期', 'CPU 费用', '内存费用', '小计'], [
        ['每日', '1.65 元', '1.92 元', '3.57 元'],
        ['每月', '49.50 元', '57.60 元', '107.10 元'],
    ])

    doc.add_heading('5.2 MySQL 数据库费用', level=2)
    add_table(doc, ['时段', '时长', '算力', '存储'], [
        ['白天活跃', '16h', '0.5 CCU', '5 GB'],
        ['深夜低活', '8h', '0.25 CCU', '5 GB'],
        ['日合计', '', '10 CCU·时', '120 GB·时'],
    ])
    add_table(doc, ['周期', '算力费用', '存储费用', '小计'], [
        ['每日', '3.42 元', '0.58 元', '4.00 元'],
        ['每月', '102.60 元', '17.46 元', '120.06 元'],
    ])

    doc.add_heading('5.3 COS 对象存储费用', level=2)
    add_table(doc, ['项目', '计算', '月费用'], [
        ['存储容量（5GB）', '5 × 30 × 0.0043', '0.65 元'],
        ['下载次数', '10 × 0.01', '0.10 元'],
        ['CDN 流量（~20GB/月）', '20 × 0.18', '3.60 元'],
        ['小计', '', '4.35 元'],
    ])

    # 六、月度总费用汇总
    doc.add_heading('六、月度总费用汇总', level=1)

    doc.add_heading('6.1 正式运营期（3 个月后）', level=2)
    add_table(doc, ['费用项', '月费用（元）', '占比'], [
        ['容器（服务）', '107.10', '43.5%'],
        ['MySQL 数据库', '120.06', '48.8%'],
        ['COS 对象存储', '4.35', '1.8%'],
        ['公网流量', '0.72', '0.3%'],
        ['构建', '2.50', '1.0%'],
        ['合计', '234.73', '100%'],
    ])

    doc.add_heading('6.2 首 3 个月（含免费额度）', level=2)
    doc.add_paragraph('首 3 个月实际费用：约 28 元/月')
    doc.add_paragraph('（容器超额 21.42 元 + COS 4.35 元 + 构建 2.50 元，MySQL 和流量被免费额度覆盖）')

    # 七、12 个月预算规划
    doc.add_heading('七、12 个月预算规划', level=1)
    add_table(doc, ['阶段', '时间', 'DAU', '月费用', '阶段费用'], [
        ['冷启动期', '第 1-3 个月', '500', '~28 元', '84 元'],
        ['稳定运营期', '第 4-6 个月', '1,000', '~235 元', '705 元'],
        ['增长期', '第 7-9 个月', '2,000', '~470 元', '1,410 元'],
        ['成熟期', '第 10-12 个月', '3,000', '~700 元', '2,100 元'],
        ['年度合计', '', '', '', '4,299 元'],
    ])

    # 八、不同规模成本对比
    doc.add_heading('八、不同规模成本对比', level=1)
    add_table(doc, ['DAU', '容器费用', '数据库费用', 'COS 费用', '月总计'], [
        ['500', '54 元', '60 元', '3 元', '117 元'],
        ['1,000', '107 元', '120 元', '4 元', '235 元'],
        ['2,000', '214 元', '240 元', '8 元', '470 元'],
        ['5,000', '536 元', '600 元', '20 元', '1,196 元'],
        ['10,000', '1,071 元', '1,201 元', '40 元', '2,392 元'],
    ])

    # 九、与传统服务器方案对比
    doc.add_heading('九、与传统服务器方案对比（1000 DAU）', level=1)
    add_table(doc, ['方案', '月费用', '优势', '劣势'], [
        ['微信云托管（推荐）', '~235 元', '自动扩缩容、免运维', '按量付费'],
        ['腾讯云 ECS + RDS', '~350 元', '固定成本可预测', '需自行运维'],
        ['轻量服务器 + 自建', '~150 元', '成本最低', '单点故障风险'],
    ])
    doc.add_paragraph('结论：微信云托管在 1000 DAU 规模下，成本低于传统 ECS 方案约 33%，且具备自动扩缩容和免运维优势。')

    # 十、投资人沟通要点
    doc.add_heading('十、投资人沟通要点', level=1)

    doc.add_heading('核心结论', level=2)
    doc.add_paragraph('1. 极低的启动成本：首 3 个月仅需约 28 元/月，年运营成本约 4,300 元')
    doc.add_paragraph('2. 弹性可扩展：从 100 到 10,000 DAU，无需迁移架构，自动扩缩容')
    doc.add_paragraph('3. 免运维：无需服务器运维人员，团队专注业务开发')
    doc.add_paragraph('4. 微信生态深度集成：天然支持小程序登录、支付、消息推送')

    doc.add_heading('建议配置路径', level=2)
    add_table(doc, ['阶段', 'DAU', '推荐配置', '月预算'], [
        ['MVP 验证期', '< 500', '0.5核1G + MySQL 0.25CCU', '100 元'],
        ['初创运营期', '500-2000', '1核2G + MySQL 0.5CCU', '300 元'],
        ['规模增长期', '2000-5000', '2核4G + MySQL 1CCU', '800 元'],
        ['成熟稳定期', '5000-10000', '4核8G + MySQL 2CCU', '2,000 元'],
    ])

    doc.add_heading('风险提示', level=2)
    doc.add_paragraph('1. 流量突增风险：爆款传播时公网流量费用可能短期上升')
    doc.add_paragraph('2. 数据库存储增长：长期运营后存储费用持续增长，建议定期清理过期数据')
    doc.add_paragraph('3. 免费额度到期：第 4 个月起费用从 ~28 元跳至 ~235 元')

    # 页脚
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— 本报告由 PlayClub 技术团队编制，仅供投资决策参考 —')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('数据来源：微信云托管官方定价 + 腾讯云 COS 定价')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 保存
    output_path = r'D:\CT电竞\PlayClub\PlayClub云托管运营成本预算报告.docx'
    doc.save(output_path)
    print(f'DOCX generated: {output_path}')

if __name__ == '__main__':
    main()
